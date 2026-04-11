from __future__ import annotations

import html
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from config import Settings, get_settings


class ReportRenderer:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()
        template_dir = self.settings.base_dir / "renderers" / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=select_autoescape(enabled_extensions=("html", "xml")),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    @staticmethod
    def _slugify(text: str) -> str:
        value = re.sub(r"[^A-Za-z0-9_-]+", "_", text)
        value = value.strip("_")
        return value or "report"

    @staticmethod
    def _normalize_paragraph_line(line: str) -> str:
        text = line.strip()
        text = re.sub(r"^[-*+]\s+", "", text)
        text = re.sub(r"^\d+[\.\)、]\s+", "", text)
        text = re.sub(r"^>\s*", "", text)
        return text

    @staticmethod
    def _set_run_font(run: Any, font_name: str, size_pt: float, bold: bool = False) -> None:
        western_font = "Times New Roman"
        run.font.name = western_font
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), western_font)
        r_fonts.set(qn("w:hAnsi"), western_font)
        r_fonts.set(qn("w:cs"), western_font)
        r_fonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size_pt)
        run.bold = bold

    @staticmethod
    def _set_paragraph_zero_spacing(paragraph: Any) -> None:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    @staticmethod
    def _replace_docx_inline_code(text: str) -> str:
        replaced = re.sub(r"`([^`]+)`", r"【\1】", text)
        return replaced.replace("`", "")

    def _parse_docx_inline_runs(self, text: str) -> list[tuple[str, bool]]:
        source = text or ""
        chunks: list[tuple[str, bool]] = []
        cursor = 0

        for match in re.finditer(r"\*\*(.+?)\*\*", source):
            plain = source[cursor : match.start()]
            if plain:
                cleaned_plain = self._replace_docx_inline_code(plain).replace("**", "")
                if cleaned_plain:
                    chunks.append((cleaned_plain, False))

            bold_text = match.group(1)
            cleaned_bold = self._replace_docx_inline_code(bold_text).replace("**", "")
            if cleaned_bold:
                chunks.append((cleaned_bold, True))

            cursor = match.end()

        tail = source[cursor:]
        if tail:
            cleaned_tail = self._replace_docx_inline_code(tail).replace("**", "")
            if cleaned_tail:
                chunks.append((cleaned_tail, False))

        if not chunks:
            cleaned = self._replace_docx_inline_code(source).replace("**", "")
            if cleaned:
                chunks.append((cleaned, False))

        return chunks

    def _add_docx_inline_runs(
        self,
        paragraph: Any,
        text: str,
        font_name: str,
        size_pt: float,
    ) -> None:
        runs = self._parse_docx_inline_runs(text)
        if not runs:
            run = paragraph.add_run("")
            self._set_run_font(run, font_name=font_name, size_pt=size_pt, bold=False)
            return

        for run_text, bold in runs:
            run = paragraph.add_run(run_text)
            self._set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold)

    def _add_docx_blank_line(
        self,
        document: Document,
        font_name: str,
        size_pt: float,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> None:
        paragraph = document.add_paragraph()
        self._set_paragraph_zero_spacing(paragraph)
        paragraph.alignment = align
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        run = paragraph.add_run(" ")
        self._set_run_font(run, font_name=font_name, size_pt=size_pt, bold=False)

    def _add_docx_cover_subtitle_line(
        self,
        document: Document,
        text: str,
        font_name: str = "楷体_GB2312",
        size_pt: float = 16,
    ) -> None:
        paragraph = document.add_paragraph()
        self._set_paragraph_zero_spacing(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        run = paragraph.add_run(text)
        self._set_run_font(run, font_name=font_name, size_pt=size_pt, bold=False)

    def _render_markdown_text(self, report_payload: dict[str, Any]) -> str:
        template = self.env.get_template("report_template.md")
        text = template.render(**report_payload)
        return text.strip() + "\n"

    def _inline_markdown_to_html(self, text: str) -> str:
        escaped = html.escape(text)
        escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
        escaped = re.sub(r"\*(.+?)\*", r"<em>\1</em>", escaped)
        escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
        return escaped

    def _markdown_to_html(self, markdown_text: str) -> str:
        blocks: list[str] = []
        paragraph_buffer: list[str] = []
        quote_buffer: list[str] = []

        def flush_paragraph() -> None:
            if not paragraph_buffer:
                return
            text = "".join(paragraph_buffer).strip()
            paragraph_buffer.clear()
            if text:
                blocks.append(f"<p>{self._inline_markdown_to_html(text)}</p>")

        def flush_quote() -> None:
            if not quote_buffer:
                return
            text = "".join(quote_buffer).strip()
            quote_buffer.clear()
            if text:
                blocks.append(
                    f"<blockquote><p>{self._inline_markdown_to_html(text)}</p></blockquote>"
                )

        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                flush_paragraph()
                flush_quote()
                continue

            if line.startswith("```"):
                continue

            if line.startswith("### "):
                flush_paragraph()
                flush_quote()
                blocks.append(f"<h3>{self._inline_markdown_to_html(line[4:].strip())}</h3>")
                continue

            if line.startswith("## "):
                flush_paragraph()
                flush_quote()
                blocks.append(f"<h2>{self._inline_markdown_to_html(line[3:].strip())}</h2>")
                continue

            if line.startswith("# "):
                flush_paragraph()
                flush_quote()
                blocks.append(f"<h2>{self._inline_markdown_to_html(line[2:].strip())}</h2>")
                continue

            if line.startswith(">"):
                flush_paragraph()
                quote_buffer.append(self._normalize_paragraph_line(line))
                continue

            flush_quote()
            paragraph_buffer.append(self._normalize_paragraph_line(line))

        flush_paragraph()
        flush_quote()
        return "\n".join(blocks)

    def _render_html_text(self, report_payload: dict[str, Any]) -> str:
        template = self.env.get_template("report_template.html")
        project_profile_html = self._markdown_to_html(
            report_payload.get("project_profile_markdown", "")
        )
        project_profile_html = re.sub(
            r"^\s*<h2>项目技术名片</h2>\s*",
            "",
            project_profile_html,
            count=1,
        )
        body_html = self._markdown_to_html(report_payload.get("body_markdown", ""))
        text = template.render(
            **report_payload,
            project_profile_html=project_profile_html,
            body_html=body_html,
        )
        return text.strip() + "\n"

    def _iter_markdown_blocks(self, markdown_text: str) -> list[tuple[str, str]]:
        blocks: list[tuple[str, str]] = []
        paragraph_buffer: list[str] = []

        def flush_paragraph() -> None:
            if not paragraph_buffer:
                return
            text = "".join(paragraph_buffer).strip()
            paragraph_buffer.clear()
            if text:
                blocks.append(("p", text))

        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                flush_paragraph()
                continue

            if line.startswith("```"):
                continue

            if line.startswith("### "):
                flush_paragraph()
                blocks.append(("h3", line[4:].strip()))
                continue

            if line.startswith("## "):
                flush_paragraph()
                blocks.append(("h2", line[3:].strip()))
                continue

            if line.startswith("# "):
                flush_paragraph()
                blocks.append(("h1", line[2:].strip()))
                continue

            paragraph_buffer.append(self._normalize_paragraph_line(line))

        flush_paragraph()
        return blocks

    def _add_docx_body_paragraph(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        self._set_paragraph_zero_spacing(paragraph)
        paragraph.paragraph_format.first_line_indent = Pt(32)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        self._add_docx_inline_runs(
            paragraph,
            text,
            font_name="仿宋_GB2312",
            size_pt=16,
        )

    def _add_docx_heading_lv1(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        self._add_docx_inline_runs(
            paragraph,
            text,
            font_name="黑体",
            size_pt=16,
        )

    def _add_docx_heading_lv2(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        self._set_paragraph_zero_spacing(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        self._add_docx_inline_runs(
            paragraph,
            text,
            font_name="楷体_GB2312",
            size_pt=16,
        )

    def _render_docx(self, docx_path: Path, report_payload: dict[str, Any]) -> None:
        document = Document()

        title = report_payload.get("title", "项目分析报告")
        subtitle = report_payload.get("subtitle", "")
        analysis_line = f"分析方向：{report_payload.get('analysis_label', '')}"
        repo_line = f"仓库地址：{report_payload.get('repo_url', '')}"
        time_line = f"生成时间：{report_payload.get('generated_at', '')}"

        title_paragraph = document.add_paragraph()
        self._set_paragraph_zero_spacing(title_paragraph)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        self._set_run_font(title_run, font_name="方正小标宋简体", size_pt=22, bold=False)

        subtitle_lines: list[str] = []
        if subtitle:
            subtitle_lines.append(subtitle)
        subtitle_lines.extend([analysis_line, repo_line, time_line])

        if subtitle_lines:
            self._add_docx_blank_line(
                document,
                font_name="楷体_GB2312",
                size_pt=16,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            for line in subtitle_lines:
                self._add_docx_cover_subtitle_line(document, line)
            self._add_docx_blank_line(
                document,
                font_name="楷体_GB2312",
                size_pt=16,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

        for markdown in (
            report_payload.get("project_profile_markdown", ""),
            report_payload.get("body_markdown", ""),
        ):
            blocks = self._iter_markdown_blocks(markdown)
            for block_type, text in blocks:
                if not text:
                    continue
                if block_type in {"h1", "h2"}:
                    self._add_docx_heading_lv1(document, text)
                elif block_type == "h3":
                    self._add_docx_heading_lv2(document, text)
                else:
                    self._add_docx_body_paragraph(document, text)

        document.save(str(docx_path))

    def render(self, report_payload: dict[str, Any]) -> dict[str, Path]:
        now = datetime.now()
        stamp = now.strftime("%Y%m%d_%H%M%S")
        repo_slug = self._slugify(report_payload.get("repo_full_name", "repo"))
        analysis_type = report_payload.get("analysis_type", "analysis")
        file_stem = f"{stamp}_{repo_slug}_{analysis_type}_report"

        md_path = self.settings.reports_output_dir / f"{file_stem}.md"
        html_path = self.settings.reports_output_dir / f"{file_stem}.html"
        docx_path = self.settings.reports_output_dir / f"{file_stem}.docx"

        md_text = self._render_markdown_text(report_payload)
        md_path.write_text(md_text, encoding="utf-8")

        html_text = self._render_html_text(report_payload)
        html_path.write_text(html_text, encoding="utf-8")

        self._render_docx(docx_path=docx_path, report_payload=report_payload)

        return {
            "md": md_path,
            "html": html_path,
            "docx": docx_path,
        }
